import io
import json
from http.server import BaseHTTPRequestHandler
from docx import Document

from http.server import BaseHTTPRequestHandler
from urllib.parse import urlparse, parse_qs
from io import BytesIO
from reportlab.pdfgen import canvas
from docx import Document

class handler(BaseHTTPRequestHandler):
    def do_GET(self):
        try:
            # Parse query params: ?type=pdf or ?type=docx
            query = parse_qs(urlparse(self.path).query)
            file_type = query.get("type", ["pdf"])[0].lower()

            if file_type == "pdf":
                buffer = BytesIO()
                p = canvas.Canvas(buffer)
                p.drawString(100, 750, "Hello, this is your resume!")
                p.showPage()
                p.save()
                data = buffer.getvalue()
                content_type = "application/pdf"
                filename = "resume.pdf"

            elif file_type == "docx":
                doc = Document()
                doc.add_heading("Resume", 0)
                doc.add_paragraph("Hello, this is your resume in DOCX format!")
                buffer = BytesIO()
                doc.save(buffer)
                data = buffer.getvalue()
                content_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                filename = "resume.docx"

            else:
                self.send_response(400)
                self.end_headers()
                self.wfile.write(b"Invalid file type")
                return

            # Send file to browser
            self.send_response(200)
            self.send_header("Content-Type", content_type)
            self.send_header("Content-Disposition", f'attachment; filename="{filename}"')
            self.end_headers()
            self.wfile.write(data)

        except Exception as e:
            self.send_response(500)
            self.end_headers()
            self.wfile.write(str(e).encode())

# This is the Vercel Serverless Function Handler
class handler(BaseHTTPRequestHandler):
    def do_POST(self):
        try:
            content_length = int(self.headers['Content-Length'])
            post_data = self.rfile.read(content_length)
            data = json.loads(post_data)

            # --- Create a new Word document in memory ---
            doc = Document()
            
            # --- Helper function for safe string conversion ---
            def s(text): return str(text) if text is not None else ""
            
            # --- Populate the Document ---
            doc.add_heading(s(data.get('name')), 0)
            p = doc.add_paragraph()
            p.add_run(s(data.get('email'))).bold = True
            p.add_run(' | ')
            p.add_run(s(data.get('phone'))).bold = True
            p.add_run(' | ')
            p.add_run(s(data.get('linkedin')))

            if data.get('summary'):
                doc.add_heading('Professional Summary', level=1)
                doc.add_paragraph(s(data.get('summary')))

            if data.get('workExperiences'):
                doc.add_heading('Work Experience', level=1)
                for job in data.get('workExperiences', []):
                    if job.get('jobTitle'):
                        doc.add_paragraph(f"{s(job.get('jobTitle'))} at {s(job.get('company'))}", style='Intense Quote')
                        date_to = 'Present' if job.get('isPresent') else s(job.get('dateTo'))
                        doc.add_paragraph(f"{s(job.get('dateFrom'))} - {date_to}")
                        for line in s(job.get('jobDescription')).split('\n'):
                            if line.strip():
                                doc.add_paragraph(line.strip(), style='List Bullet')
            
            if data.get('educations'):
                doc.add_heading('Education', level=1)
                for edu in data.get('educations', []):
                     if edu.get('degree'):
                        doc.add_paragraph(f"{s(edu.get('degree'))} at {s(edu.get('school'))}", style='Intense Quote')
                        date_to = 'Present' if edu.get('isPresent') else s(edu.get('dateTo'))
                        doc.add_paragraph(f"{s(edu.get('dateFrom'))} - {date_to}")
            
            if data.get('skills'):
                doc.add_heading('Skills', level=1)
                skills_list = [skill.strip() for skill in s(data.get('skills')).split(',') if skill.strip()]
                for skill in skills_list:
                    doc.add_paragraph(skill, style='List Bullet')
            
            # --- Save the document to an in-memory stream ---
            file_stream = io.BytesIO()
            doc.save(file_stream)
            file_stream.seek(0)
            doc_bytes = file_stream.read()

            # --- Send the file back to the user ---
            self.send_response(200)
            self.send_header('Content-Type', 'application/vnd.openxmlformats-officedocument.wordprocessingml.document')
            self.send_header('Content-Disposition', f'attachment; filename="Resume - {s(data.get("name"))}.docx"')
            self.end_headers()
            self.wfile.write(doc_bytes)

        except Exception as e:
            print(f"MAIN ERROR: {e}")
            self.send_response(500)
            self.send_header('Content-Type', 'application/json')
            self.end_headers()
            error_response = json.dumps({'error': 'An internal server error occurred.'})
            self.wfile.write(error_response.encode('utf-8'))
        
        return