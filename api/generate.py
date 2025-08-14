import io
import json
from http.server import BaseHTTPRequestHandler
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.lib.colors import HexColor

class handler(BaseHTTPRequestHandler):
    def do_POST(self):
        try:
            content_length = int(self.headers['Content-Length'])
            post_data = self.rfile.read(content_length)
            data = json.loads(post_data)

            file_format = data.get('format', 'docx')
            user_name = data.get('name', 'resume')
            accent_color_hex = data.get('accentColor', '#0d47a1').lstrip('#')
            
            content_type = ""
            download_name = ""
            file_bytes = None

            def s(text): return str(text) if text is not None else ""

            # --- GENERATE PDF ---
            if file_format == 'pdf':
                content_type = 'application/pdf'
                download_name = f"Resume - {s(user_name)}.pdf"
                
                buffer = io.BytesIO()
                doc = SimpleDocTemplate(buffer, topMargin=0.5*inch, bottomMargin=0.5*inch, leftMargin=0.75*inch, rightMargin=0.75*inch)
                
                # --- PDF Styles ---
                styles = getSampleStyleSheet()
                accent_color = HexColor(f"#{accent_color_hex}")
                styles.add(ParagraphStyle(name='Name', fontName='Helvetica-Bold', fontSize=24, spaceAfter=2))
                styles.add(ParagraphStyle(name='Contact', fontName='Helvetica', fontSize=10, spaceAfter=12))
                styles.add(ParagraphStyle(name='SectionHead', fontName='Helvetica-Bold', fontSize=14, spaceBefore=6, spaceAfter=6, textColor=accent_color))
                styles.add(ParagraphStyle(name='JobTitle', fontName='Helvetica-Bold', fontSize=11, spaceAfter=2))
                styles.add(ParagraphStyle(name='JobDate', fontName='Helvetica-Oblique', fontSize=9, spaceAfter=4))
                styles.add(ParagraphStyle(name='Bullet', fontName='Helvetica', fontSize=10, leftIndent=18, spaceAfter=2))
                
                story = []

                # --- Populate PDF ---
                story.append(Paragraph(s(data.get('name')), styles['Name']))
                story.append(Paragraph(f"{s(data.get('email'))} | {s(data.get('phone'))} | {s(data.get('linkedin'))}", styles['Contact']))
                
                if data.get('summary'):
                    story.append(Paragraph("Summary", styles['SectionHead']))
                    story.append(Paragraph(s(data.get('summary')), styles['BodyText']))
                
                if data.get('workExperiences'):
                    story.append(Paragraph("Work Experience", styles['SectionHead']))
                    for job in data.get('workExperiences', []):
                        if job.get('jobTitle'):
                            story.append(Paragraph(f"{s(job.get('jobTitle'))} | {s(job.get('company'))}", styles['JobTitle']))
                            date_to = 'Present' if job.get('isPresent') else s(job.get('dateTo'))
                            story.append(Paragraph(f"{s(job.get('dateFrom'))} - {date_to}", styles['JobDate']))
                            for line in s(job.get('jobDescription')).split('\n'):
                                if line.strip(): story.append(Paragraph(f"• {line.strip()}", styles['Bullet']))
                            story.append(Spacer(1, 0.1*inch))
                
                if data.get('educations'):
                    story.append(Paragraph("Education", styles['SectionHead']))
                    for edu in data.get('educations', []):
                        if edu.get('degree'):
                            story.append(Paragraph(f"{s(edu.get('degree'))} | {s(edu.get('school'))}", styles['JobTitle']))
                            date_to = 'Present' if edu.get('isPresent') else s(edu.get('dateTo'))
                            story.append(Paragraph(f"{s(edu.get('dateFrom'))} - {date_to}", styles['JobDate']))

                if data.get('skills'):
                    story.append(Paragraph("Skills", styles['SectionHead']))
                    skills_list = [skill.strip() for skill in s(data.get('skills')).split(',') if skill.strip()]
                    for skill in skills_list:
                        story.append(Paragraph(f"• {skill}", styles['Bullet']))

                doc.build(story)
                file_bytes = buffer.getvalue()

            # --- GENERATE DOCX ---
            else:
                content_type = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
                download_name = f"Resume - {s(user_name)}.docx"
                doc = Document()

                # --- DOCX Styles ---
                rgb = tuple(int(accent_color_hex[i:i+2], 16) for i in (0, 2, 4))
                
                # --- Populate DOCX ---
                name_p = doc.add_paragraph()
                name_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                runner = name_p.add_run(s(data.get('name')))
                runner.font.name = 'Calibri'
                runner.font.size = Pt(24)
                runner.font.bold = True

                contact_p = doc.add_paragraph()
                contact_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                contact_p.add_run(f"{s(data.get('email'))} | {s(data.get('phone'))} | {s(data.get('linkedin'))}")

                if data.get('summary'):
                    h = doc.add_heading("Summary", level=1)
                    h.runs[0].font.color.rgb = RGBColor(rgb[0], rgb[1], rgb[2])
                    doc.add_paragraph(s(data.get('summary')))

                if data.get('workExperiences'):
                    h = doc.add_heading("Work Experience", level=1)
                    h.runs[0].font.color.rgb = RGBColor(rgb[0], rgb[1], rgb[2])
                    for job in data.get('workExperiences', []):
                        if job.get('jobTitle'):
                            job_p = doc.add_paragraph()
                            job_p.add_run(f"{s(job.get('jobTitle'))} | {s(job.get('company'))}").bold = True
                            date_to = 'Present' if job.get('isPresent') else s(job.get('dateTo'))
                            doc.add_paragraph(f"{s(job.get('dateFrom'))} - {date_to}").paragraph_format.space_after = Pt(4)
                            for line in s(job.get('jobDescription')).split('\n'):
                                if line.strip(): doc.add_paragraph(line.strip(), style='List Bullet')
                
                if data.get('educations'):
                    h = doc.add_heading("Education", level=1)
                    h.runs[0].font.color.rgb = RGBColor(rgb[0], rgb[1], rgb[2])
                    for edu in data.get('educations', []):
                        if edu.get('degree'):
                            edu_p = doc.add_paragraph()
                            edu_p.add_run(f"{s(edu.get('degree'))} | {s(edu.get('school'))}").bold = True
                            date_to = 'Present' if edu.get('isPresent') else s(edu.get('dateTo'))
                            doc.add_paragraph(f"{s(edu.get('dateFrom'))} - {date_to}")

                if data.get('skills'):
                    h = doc.add_heading("Skills", level=1)
                    h.runs[0].font.color.rgb = RGBColor(rgb[0], rgb[1], rgb[2])
                    skills_list = [skill.strip() for skill in s(data.get('skills')).split(',') if skill.strip()]
                    for skill in skills_list:
                        doc.add_paragraph(skill, style='List Bullet')
                
                file_stream = io.BytesIO()
                doc.save(file_stream)
                file_stream.seek(0)
                file_bytes = file_stream.read()

            # --- Send the file back to the user ---
            self.send_response(200)
            self.send_header('Content-Type', content_type)
            self.send_header('Content-Disposition', f'attachment; filename="{download_name}"')
            self.end_headers()
            self.wfile.write(file_bytes)

        except Exception as e:
            print(f"MAIN ERROR: {e}")
            self.send_response(500)
            self.send_header('Content-Type', 'application/json')
            self.end_headers()
            error_response = json.dumps({'error': 'An internal server error occurred.'})
            self.wfile.write(error_response.encode('utf-8'))
        
        return